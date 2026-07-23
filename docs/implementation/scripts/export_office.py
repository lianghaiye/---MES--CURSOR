#!/usr/bin/env python3
"""Export implementation handbook Markdown to Word (.docx) and field attachments to Excel (.xlsx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "office"
MD_TABLE_SPLIT = re.compile(r"^\|[\s\-:|]+\|$")
HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
UL = re.compile(r"^(\s*)[-*]\s+(.*)$")
OL = re.compile(r"^(\s*)\d+\.\s+(.*)$")
CHECK = re.compile(r"^(\s*)-\s+\[([ xX])\]\s+(.*)$")
BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_runs(paragraph, text: str, base_size=11):
    """Split **bold**, `code`, [links](url) into runs."""
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=base_size)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(BOLD.match(token).group(1))
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(INLINE_CODE.match(token).group(1))
            set_run_font(run, size=base_size)
            run.font.name = "Consolas"
        else:
            lm = LINK.match(token)
            run = paragraph.add_run(f"{lm.group(1)} ({lm.group(2)})")
            set_run_font(run, size=base_size, color=RGBColor(0x1A, 0x56, 0xDB))
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [BOLD.sub(r"\1", c) for c in cells]


def md_to_docx(md_path: Path, docx_path: Path, title: str | None = None):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_run_font(run, size=18, bold=True)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        hm = HEADING.match(stripped)
        if hm:
            level = len(hm.group(1))
            text = BOLD.sub(r"\1", hm.group(2))
            # Skip duplicate H1 if we already set title
            if level == 1 and title and text.replace(" ", "") in title.replace(" ", ""):
                i += 1
                continue
            heading = doc.add_heading(text, level=min(level, 3))
            for run in heading.runs:
                set_run_font(run, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and MD_TABLE_SPLIT.match(lines[i + 1].strip()):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not MD_TABLE_SPLIT.match(lines[i].strip()):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    val = row[c_idx] if c_idx < len(row) else ""
                    run = p.add_run(val)
                    set_run_font(run, size=9, bold=(r_idx == 0))
            doc.add_paragraph()
            continue

        cm = CHECK.match(line)
        if cm:
            mark = "☑" if cm.group(2).lower() == "x" else "☐"
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, f"{mark} {cm.group(3)}")
            i += 1
            continue

        um = UL.match(line)
        if um:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, um.group(2))
            i += 1
            continue

        om = OL.match(line)
        if om:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, om.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            quote = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(quote)
            set_run_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Word: {docx_path.relative_to(ROOT)}")


def sheet_name(title: str, used: set[str]) -> str:
    name = re.sub(r"^Sheet\d+\s*", "", title).strip()
    name = re.sub(r"[\\/*?:\[\]]", "_", name)[:31] or "Sheet"
    base = name
    n = 1
    while name in used:
        suffix = f"_{n}"
        name = (base[: 31 - len(suffix)] + suffix)
        n += 1
    used.add(name)
    return name


def style_header(ws, col_count: int):
    fill = PatternFill("solid", fgColor="1F4E79")
    font = Font(bold=True, color="FFFFFF", name="微软雅黑", size=11)
    thin = Border(
        left=Side(style="thin", color="D0D0D0"),
        right=Side(style="thin", color="D0D0D0"),
        top=Side(style="thin", color="D0D0D0"),
        bottom=Side(style="thin", color="D0D0D0"),
    )
    for c in range(1, col_count + 1):
        cell = ws.cell(1, c)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def autosize(ws):
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        width = 12
        for cell in col:
            if cell.value:
                width = max(width, min(48, len(str(cell.value)) + 2))
        ws.column_dimensions[letter].width = width


def md_tables_to_xlsx(md_path: Path, xlsx_path: Path):
    """Each ## section that contains a markdown table becomes one Excel sheet."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    wb = Workbook()
    wb.remove(wb.active)
    used: set[str] = set()

    # Meta sheet: title + blockquotes + non-table intro
    meta = wb.create_sheet("说明", 0)
    meta["A1"] = "来源"
    meta["B1"] = md_path.name
    meta["A2"] = "说明"
    row = 3
    i = 0
    current_title = "数据"
    pending_title = None
    tables_written = 0

    while i < len(lines):
        stripped = lines[i].strip()
        hm = HEADING.match(stripped)
        if hm and len(hm.group(1)) <= 2:
            pending_title = BOLD.sub(r"\1", hm.group(2))
            current_title = pending_title
            i += 1
            continue

        if stripped.startswith(">") and tables_written == 0:
            meta.cell(row, 1, stripped.lstrip("> ").strip())
            row += 1
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and MD_TABLE_SPLIT.match(lines[i + 1].strip()):
            headers = parse_table_row(stripped)
            i += 2
            data = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not MD_TABLE_SPLIT.match(lines[i].strip()):
                    data.append(parse_table_row(lines[i]))
                i += 1
            title = pending_title or current_title or f"表{tables_written + 1}"
            ws = wb.create_sheet(sheet_name(title, used))
            for c, h in enumerate(headers, 1):
                ws.cell(1, c, h)
            for r_idx, data_row in enumerate(data, 2):
                for c_idx, val in enumerate(data_row, 1):
                    ws.cell(r_idx, c_idx, val)
                    ws.cell(r_idx, c_idx).alignment = Alignment(wrap_text=True, vertical="center")
                    ws.cell(r_idx, c_idx).font = Font(name="微软雅黑", size=10)
            # Extra blank rows for customer fill (collection sheets)
            if "确认" not in title and data:
                for extra in range(20):
                    for c in range(1, len(headers) + 1):
                        ws.cell(len(data) + 2 + extra, c, "")
            style_header(ws, len(headers))
            autosize(ws)
            tables_written += 1
            pending_title = None
            continue

        # Capture signature / checklist lines onto 说明
        if stripped.startswith("**") and tables_written == 0:
            meta.cell(row, 1, BOLD.sub(r"\1", stripped))
            row += 1
        i += 1

    if tables_written == 0:
        # Fallback: whole file as one sheet of paragraphs
        ws = wb.create_sheet("内容")
        ws["A1"] = "内容"
        r = 2
        for line in lines:
            if line.strip() and not line.strip().startswith("#"):
                ws.cell(r, 1, BOLD.sub(r"\1", line.strip()))
                r += 1
        style_header(ws, 1)
        autosize(ws)

    style_header(meta, 2)
    autosize(meta)

    xlsx_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(xlsx_path)
    print(f"Excel: {xlsx_path.relative_to(ROOT)} ({tables_written} sheets)")


def main():
    OUT.mkdir(parents=True, exist_ok=True)

    md_to_docx(
        ROOT / "I-DOMS场景化实施手册.md",
        OUT / "I-DOMS场景化实施手册.docx",
        title="I-DOMS 场景化实施手册",
    )

    for card in sorted((ROOT / "scene-cards").glob("*.md")):
        md_to_docx(card, OUT / "scene-cards" / f"{card.stem}.docx", title=card.stem)

    for cl in sorted((ROOT / "checklists").glob("*.md")):
        md_to_docx(cl, OUT / "checklists" / f"{cl.stem}.docx", title=cl.stem)

    for att in sorted((ROOT / "attachments").glob("*.md")):
        md_tables_to_xlsx(att, OUT / "attachments" / f"{att.stem}.xlsx")

    print(f"\nDone. Output directory: {OUT}")


if __name__ == "__main__":
    main()
