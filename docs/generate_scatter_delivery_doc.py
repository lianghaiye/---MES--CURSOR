#!/usr/bin/env python3
"""Generate Word doc: I-DOMS 散件发货功能说明与规则"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date

OUTPUT = "/Users/jianjianya/CURSOR/i-doms-web/docs/I-DOMS-散件发货功能说明与规则.docx"


def set_cell_shading(cell, color="E7F3FF"):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    return h


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, indent=0.75):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E8FB")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "PingFang SC"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "PingFang SC"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Cover title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("I-DOMS 散件发货\n功能说明与业务规则")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
    tr.font.name = "PingFang SC"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"文档版本：V1.0    生成日期：{date.today().isoformat()}\n基于 i-doms-web 当前实现整理")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sr.font.name = "PingFang SC"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    doc.add_paragraph()

    # 1. 概述
    add_heading(doc, "一、功能概述", 1)
    add_para(
        doc,
        "散件发货是 I-DOMS 销售发货能力中的一种交付方式。与「整机发货」按产品数量整批出库不同，"
        "散件发货面向交付方式为「散件」的销售订单明细行，允许用户按 EBOM（工程物料清单）"
        "逐物料勾选本次发运数量，实现零部件、组件的分批发运。",
    )
    add_para(doc, "适用场景包括但不限于：", bold=True)
    add_bullet(doc, "客户分阶段收货，先发关键零部件、后发其余物料")
    add_bullet(doc, "软件/零部件类履约（历史订单头「履约方式」为「软件零部件发货」时，行级默认映射为散件）")
    add_bullet(doc, "同一销售订单中整机行与散件行并存，可在一次「申请发货」中分别处理")

    add_para(doc, "功能入口：销售管理 → 销售订单 → 勾选订单 → 申请发货。", bold=True)
    add_para(
        doc,
        "在申请发货弹窗中，系统按销售明细行的交付方式自动拆分为「整机发货」与「散件发运」两个区块；"
        "散件行通过「选择发运物料」抽屉从 EBOM 中拣选物料。",
    )

    # 2. 概念与数据
    add_heading(doc, "二、核心概念与数据来源", 1)

    add_heading(doc, "2.1 交付方式（行级）", 2)
    add_table(
        doc,
        ["取值", "说明"],
        [
            ["整机", "按产品数量发货，填写「本次发货数量」"],
            ["散件", "按 EBOM 物料拣选发运，不填写产品级发货数量"],
        ],
    )
    add_para(doc, "交付方式判定规则：")
    add_bullet(doc, "销售订单明细行若已设置 deliveryMode 为「整机」或「散件」，以行级为准")
    add_bullet(doc, "若行级未设置，且订单头履约方式为「软件零部件发货」，则映射为「散件」")
    add_bullet(doc, "其余情况默认为「整机」")

    add_heading(doc, "2.2 EBOM 快照", 2)
    add_para(
        doc,
        "散件发运依赖销售订单明细行上的 EBOM 快照（ebomSnapshot）。"
        "自产销售订单审核通过时，系统为每条明细行：",
    )
    add_bullet(doc, "校验产品已关联、且存在「使用中」状态的产品 BOM")
    add_bullet(doc, "按销售数量 × BOM 结构展开物料树，生成 ebomSnapshot（含 materials 树形结构）")
    add_bullet(doc, "同步规范化行级 deliveryMode")
    add_para(
        doc,
        "若 EBOM 快照为空（未审核、无 BOM、或物料树未生成），散件发运抽屉将提示「无 EBOM 物料，"
        "请确认订单已审核并生成 EBOM」。",
    )

    add_heading(doc, "2.3 散件发运数据结构", 2)
    add_table(
        doc,
        ["字段/对象", "说明"],
        [
            ["scatterShipments", "申请发货表单中的散件发运行列表，每行对应一条散件销售明细"],
            ["materialPicks", "由 EBOM 展平生成的可拣选物料行（含需求、库存、缺口、发运数量等）"],
            ["ebomSnapshotId", "关联的 EBOM 快照 ID"],
            ["remark", "散件行级发运备注（在选择发运物料抽屉中填写）"],
        ],
    )

    # 3. 功能描述
    add_heading(doc, "三、功能描述", 1)

    add_heading(doc, "3.1 销售订单维护（前置）", 2)
    add_bullet(doc, "新建/编辑销售订单时，明细行可设置「交付方式」为整机或散件（默认整机）")
    add_bullet(doc, "自产销售订单审核通过后自动绑定 BOM 并生成 EBOM 快照，为散件发运提供物料清单")
    add_bullet(doc, "生产计划「工作项」列表展示各行交付方式标签（散件为橙色、整机为蓝色）")

    add_heading(doc, "3.2 申请发货弹窗 — 散件发运区块", 2)
    add_para(doc, "打开申请发货时，系统自动：")
    add_bullet(doc, "筛选 deliveryMode = 散件 的销售明细，初始化 scatterShipments")
    add_bullet(doc, "从 ebomSnapshot（或 line.materials）构建 materialPicks 拣选列表")
    add_bullet(doc, "计算并展示各行发货状态、发货进度")

    add_para(doc, "散件发运产品行展示字段（与整机基本一致，不含「本次发货数量」列）：", bold=True)
    add_bullet(doc, "产品名称、编码、规格、订单数量、单价、交付方式、发货状态、发货进度等")
    add_bullet(doc, "可编辑：本次发货单价（不含税）、行备注")
    add_bullet(doc, "操作：「选择发运物料」")

    add_para(doc, "展开行（已选发运物料）：", bold=True)
    add_bullet(doc, "展示已勾选且发运数量 > 0 的 EBOM 物料明细")
    add_bullet(doc, "支持单行删除已选物料")
    add_bullet(doc, "展示发运备注")

    add_heading(doc, "3.3 选择发运物料抽屉", 2)
    add_para(doc, "点击「选择发运物料」打开侧滑抽屉，主要能力：")
    add_bullet(doc, "展示 EBOM 展平物料树（子级物料缩进显示）")
    add_bullet(doc, "列：勾选、物料名称、编码、需求、可用库存、缺口、单位、本次发运数量")
    add_bullet(doc, "缺口 > 0 时以红色高亮提示")
    add_bullet(doc, "勾选物料时，若发运数量为 0，自动填入需求数量（至少为 1）")
    add_bullet(doc, "取消勾选时，发运数量清零")
    add_bullet(doc, "可填写发运备注（选填）")

    add_heading(doc, "3.4 整机与散件混合发运", 2)
    add_para(
        doc,
        "同一销售订单可同时包含整机行与散件行。申请发货弹窗分别展示「整机发货」与「散件发运」两个表格；"
        "提交时两类明细均须满足各自校验规则。若两类均无明细，则不可提交。",
    )

    add_heading(doc, "3.5 发货进度与状态计算", 2)
    add_table(
        doc,
        ["指标", "计算公式 / 规则"],
        [
            ["散件发货进度", "本次已选发运数量合计 / EBOM 全部物料需求数量之和"],
            ["散件行发货状态 — 未发货", "无任何已勾选且发运数量 > 0 的物料"],
            ["散件行发货状态 — 部分发货", "有已选物料，且发运合计 < EBOM 需求合计"],
            ["散件行发货状态 — 已发完", "发运合计 ≥ EBOM 需求合计"],
            ["整机发货进度", "已发货数量 / 订单数量（整数展示）"],
        ],
    )

    add_table(
        doc,
        ["物料拣选行字段", "计算规则"],
        [
            ["需求数量 demandQty", "来自 EBOM 快照展开结果"],
            ["可用库存 availableStock", "优先取 availableStock，否则取 stockQty"],
            ["缺口 gapQty", "max(0, 需求数量 − 可用库存)"],
            ["发货总额 deliveryAmountExTax", "本次发货单价（不含税）× 对应计价基数（散件行按产品行单价字段维护）"],
        ],
    )

    # 4. 业务规则
    add_heading(doc, "四、业务规则", 1)

    add_heading(doc, "4.1 前置条件", 2)
    add_bullet(doc, "须在销售订单列表勾选一条订单后，方可打开「申请发货」")
    add_bullet(doc, "散件发运行须对应 deliveryMode = 散件 的销售明细")
    add_bullet(doc, "自产订单须已审核，且明细行已生成 EBOM 快照，否则无法拣选物料")

    add_heading(doc, "4.2 申请发货 — 公共校验（整机 + 散件共用）", 2)
    add_table(
        doc,
        ["校验项", "规则"],
        [
            ["发货编码", "必填，可点击「生成编码」自动生成（格式 SH + 日期 + 流水）"],
            ["客户名称", "必填"],
            ["发货方式", "必填（如物流、自提等）"],
            ["出库仓库", "当「申请出库」开关开启时必填"],
            ["可发运明细", "整机区块与散件区块至少有一类存在明细，否则提示「本单无整机或散件可发运明细」"],
        ],
    )

    add_heading(doc, "4.3 散件发运 — 提交校验", 2)
    add_table(
        doc,
        ["校验项", "规则", "提示信息"],
        [
            ["每行须选物料", "每条 scatterShipment 至少有一条 materialPick：selected=true 且 shipQty>0", "散件行「{产品名}」请选择发运物料"],
        ],
    )
    add_para(
        doc,
        "说明：散件发运不要求单次发运必须覆盖全部 EBOM 需求；"
        "允许部分物料、部分数量发运，系统据此更新「部分发货」状态。",
    )

    add_heading(doc, "4.4 选择发运物料抽屉 — 校验", 2)
    add_table(
        doc,
        ["校验项", "规则", "提示信息"],
        [
            ["至少选一项", "须至少勾选一条 EBOM 物料且发运数量 > 0", "请至少勾选一项 EBOM 物料，并填写本次发运数量"],
            ["勾选须填数量", "已勾选行的 shipQty 须 > 0", "物料「{名称}」已勾选，请填写发运数量"],
        ],
    )

    add_heading(doc, "4.5 整机发货校验（对比参考）", 2)
    add_para(doc, "与散件同属申请发货弹窗，规则如下（散件行不适用「本次发货数量」列）：")
    add_table(
        doc,
        ["校验项", "规则"],
        [
            ["本次发货数量", "必填"],
            ["数量范围", "须 > 0，且 ≤ 可发数量（订单数量 − 已发数量）"],
        ],
    )

    add_heading(doc, "4.6 物料拣选交互规则", 2)
    add_bullet(doc, "勾选物料：若发运数量为 0，默认填入 max(需求数量, 1)")
    add_bullet(doc, "取消勾选：发运数量置 0，并从「已选发运物料」列表移除")
    add_bullet(doc, "在展开区删除已选物料：等效于取消勾选并重算发货状态")
    add_bullet(doc, "保存抽屉后：自动刷新行发货状态；有已选物料时自动展开该行")

    add_heading(doc, "4.7 与出库 / 质检链路", 2)
    add_para(
        doc,
        "申请发货弹窗支持「申请出库」开关及出库仓库选择。"
        "当前前端实现中，确认提交后触发 confirmed 事件并提示「发货申请已提交」；"
        "后续与出库单、出厂质检的自动联动以服务端接口对接为准。",
    )
    add_bullet(doc, "非销售出库类型或未走销售出库流程时，出库/质检规则不受散件拣选界面约束")
    add_bullet(doc, "散件发运的核心差异在于出库明细粒度为 EBOM 物料级，而非产品整件数量")

    # 5. 流程
    add_heading(doc, "五、业务流程", 1)
    add_para(doc, "散件发货典型流程：", bold=True)
    steps = [
        "维护销售订单，明细行设置交付方式为「散件」",
        "审核销售订单（自产）：系统校验 BOM、生成 EBOM 快照",
        "销售订单 → 申请发货：填写发货信息",
        "在「散件发运」区块点击「选择发运物料」",
        "勾选 EBOM 物料、填写本次发运数量、可选填发运备注 → 确定",
        "确认各行已选物料无误，提交申请发货",
        "（后续）生成出库单 / 执行出库 / 必要时出厂质检",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    # 6. 界面与模块
    add_heading(doc, "六、涉及模块与代码位置", 1)
    add_table(
        doc,
        ["模块/文件", "职责"],
        [
            ["ApplyDeliveryModal.vue", "申请发货弹窗，整机/散件分区、提交校验"],
            ["ScatterShipDrawer.vue", "选择发运物料抽屉，EBOM 拣选"],
            ["utils/salesDeliveryMode.js", "交付方式枚举与行级判定"],
            ["utils/deliveryLine.js", "发货行映射、整机/散件展示行构建"],
            ["utils/shipEbom.js", "EBOM 展平、拣选行、进度/状态计算"],
            ["utils/ebomSnapshot.js", "审核时 BOM → EBOM 快照展开"],
            ["store/salesOrderStore.js", "审核通过时生成 EBOM、生产计划"],
            ["CreateSalesOrderModal.vue", "销售明细行交付方式维护"],
            ["ProductionPlanView.vue", "工作项展示交付方式标签"],
        ],
    )

    # 7. 附录
    add_heading(doc, "七、附录：状态枚举", 1)
    add_table(
        doc,
        ["类型", "枚举值"],
        [
            ["交付方式", "整机、散件"],
            ["行发货状态", "未发货、部分发货、已发完"],
            ["EBOM 物料行结果（质检场景，非本功能必填）", "合格、不合格"],
        ],
    )

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build()
