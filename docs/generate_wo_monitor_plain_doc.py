#!/usr/bin/env python3
"""生成大白话版：工单监管看板说明（Word）"""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = "/Users/jianjianya/CURSOR/i-doms-web/docs/工单监管看板-大白话说明.docx"


def set_cell_shading(cell, color="E7F3FF"):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:fill"): color,
            qn("w:val"): "clear",
        },
    )
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    return h


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, indent=0.75):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("要点：" + text)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x56, 0x06)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "D9E8FB")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "PingFang SC"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "PingFang SC"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("工单监管看板\n大白话说明")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
    tr.font.name = "PingFang SC"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"给产品 / 业务 / 联调同学看的口语版\n"
        f"生成日期：{date.today().isoformat()}    对应技术文档：工单监管看板-取数与交互说明.md"
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sr.font.name = "PingFang SC"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    doc.add_paragraph()

    # 1
    add_heading(doc, "一、这个看板是干什么的？", 1)
    add_para(
        doc,
        "车间大屏用的「工单监管看板」。一眼能看到：现在有多少单在干、工序干到哪、"
        "工人谁忙谁闲。入口一般是：更多 → 看板管理 → 工单监管看板；会新开一个全屏标签页。",
    )
    add_para(doc, "大屏上大致分三块：", bold=True)
    add_bullet(doc, "上面一排数字（KPI）：工单数量、报工数量、工序状态")
    add_bullet(doc, "中间左边：工单执行情况（每张工单的工艺路线进度）")
    add_bullet(doc, "右边：工人接单情况（小组 + 单工人）")

    # 2
    add_heading(doc, "二、数据从哪来？", 1)
    add_para(doc, "看板自己不存业务数据，是把现有几类数据拼起来看：")
    add_table(
        doc,
        ["数据", "用在哪"],
        [
            ["生产工单、组装工单", "上面的工单数字、中间工单列表、工艺路线"],
            ["工序移动任务（工人端任务）", "工序显示成啥状态、工序 KPI、工人忙不忙"],
            ["报工记录", "上面「报工数量」那几项"],
            ["员工小组", "右边小组列表"],
        ],
    )
    add_para(
        doc,
        "先把生产工单和组装工单合在一起，再按你选的「工单类型、工作中心」过滤，"
        "后面算数字、列清单，基本都围着这批工单转。",
    )

    # 3
    add_heading(doc, "三、上面几个筛选怎么理解？", 1)
    add_table(
        doc,
        ["筛选项", "大白话"],
        [
            [
                "周期（本日/本周/本月）",
                "只影响「已完成工单数」和「报工数量」。待下发、进行中看的是此刻有多少，不跟周期走。",
            ],
            ["工单类型", "只看生产、只看组装，还是都看。"],
            ["工作中心", "只看某个车间/中心的工单。"],
            [
                "列表状态（进行中/待下发/已完成/全部）",
                "只改中间工单列表显示哪些单。上面 KPI 数字不受这个切换影响。",
            ],
        ],
    )
    add_callout(
        doc,
        "很多人会以为切「本日」会把进行中的单也滤掉——不会。进行中、待下发永远是「现在这一刻」的快照。",
    )

    # 4
    add_heading(doc, "四、工单状态怎么归类？", 1)
    add_para(doc, "系统里状态很多，看板先把它们归成几大类：")
    add_table(
        doc,
        ["归类", "包含哪些状态", "说明"],
        [
            ["待下发", "待下发", "还没真正开干"],
            [
                "进行中",
                "已下发、执行中、暂停……（除待下发/完成/终止外）",
                "注意：不止「执行中」！已下发、暂停也算进行中",
            ],
            ["已完成", "已完成、完成", "算完了"],
            ["不算", "终止", "不进 KPI，也不算进行中"],
        ],
    )
    add_callout(doc, "「进行中」≠ 仅「执行中」。已下发、暂停都算进行中。联调时最容易踩这个坑。")

    # 5
    add_heading(doc, "五、上面 KPI 怎么数？", 1)

    add_heading(doc, "5.1 工单数量", 2)
    add_bullet(doc, "待下发：现在有多少单还是「待下发」（此刻快照）")
    add_bullet(doc, "进行中：现在有多少单落在「进行中」这一类（此刻快照）")
    add_bullet(doc, "已完成：在你选的周期里，完成了多少单（看更新时间，没有就看创建时间）")

    add_heading(doc, "5.2 报工数量", 2)
    add_para(doc, "看报工记录表，满足下面条件才算：")
    add_bullet(doc, "报工时间落在所选周期内")
    add_bullet(doc, "有工单号/工单 ID 的：必须能对应到当前筛出来的那批工单")
    add_bullet(
        doc,
        "没有挂工单的散报工：只有「类型=全部 且 工作中心=全部」时才计入，"
        "免得你筛了车间，还把别的散报工算进来",
    )
    add_para(doc, "三项分别是：报工笔数、良品合计、不良合计。")

    add_heading(doc, "5.3 工序状态（已报工 / 在制 / 待领取）", 2)
    add_para(
        doc,
        "先给每张工单画出工艺路线上的工序节点，再按节点状态归类加总。"
        "现网约定很重要：工人提交一次报工，这道工序任务就算完成了——"
        "不要求报工数量必须凑满排产数。",
    )
    add_table(
        doc,
        ["KPI", "大白话"],
        [
            ["已报工", "这道工序任务已经报完/完成了"],
            [
                "在制",
                "活已经到工人手里了，但还没报完。包括：执行中、待开始、待报工、待分发",
            ],
            ["待领取", "多人抢领、还没人领走的那种"],
            ["不算进上面三项", "暂停的工序"],
        ],
    )
    add_para(
        doc,
        "极简单人/单组下发后，任务经常直接变成「执行中/待开始」，没有「待领取」这一步——"
        "这种算在制，不算待领取。",
    )

    # 6
    add_heading(doc, "六、中间工单列表怎么看？", 1)

    add_heading(doc, "6.1 列表里有哪些单", 2)
    add_para(
        doc,
        "在已过滤的工单里，再按「列表状态」筛一遍，按最近更新时间倒序排，再分页。"
        "每页显示几条，会按大屏可视高度自动算，尽量不出现纵向滚动条；到时间会自动翻页轮播。",
    )

    add_heading(doc, "6.2 一行工单上显示什么", 2)
    add_bullet(doc, "单号、产品名、产品编号、规格、材质、生产/组装、状态")
    add_bullet(doc, "工作中心、计划数/排产数、计划日期")
    add_bullet(doc, "有多批次时可能显示「批次#1/2」这类标签")
    add_bullet(doc, "计划结束日早于今天、且单还没完成/终止 → 显示「逾期」")
    add_bullet(doc, "下面一条工艺路线：串行工序横着排；同一步号有多道工序 = 并行（虚线框包起来）")

    add_heading(doc, "6.3 工艺路线节点怎么定", 2)
    add_bullet(doc, "优先用工单上的工序清单；再尽量匹配工人端的移动任务来定状态和数量")
    add_bullet(doc, "没有工序清单时，就按任务里的工序名拼路线")
    add_bullet(doc, "同一「步骤号」的多道工序 = 并行；步骤号从小到大串成整条路线")
    add_para(doc, "节点上「计划 / 良 / 不良」的取值优先级（有值就用，从前到后找）：", bold=True)
    add_bullet(doc, "计划：工序计划/排产 → 任务期望/目标 → 工单排产/计划")
    add_bullet(doc, "良：任务报工良品 → 工序上的良品字段")
    add_bullet(doc, "不良：任务报工不良 → 工序上的不良字段")

    # 7
    add_heading(doc, "七、右边工人区怎么理解？", 1)

    add_heading(doc, "7.1 什么叫「忙」", 2)
    add_para(
        doc,
        "个人任务状态落在「待报工、待开始、执行中」就算忙。"
        "小组任务再多算一种「待分发」。看的是认领人/执行人姓名。",
    )

    add_heading(doc, "7.2 工人小组", 2)
    add_bullet(doc, "只展示启用中的小组")
    add_bullet(doc, "小组有进行中任务，或组里有人个人在忙 → 整组偏「忙」")
    add_bullet(
        doc,
        "组长接了小组任务，组员自己手头没有忙活 → 组员显示「组内待命」"
        "（不是空闲！人还绑在这组活上）",
    )
    add_bullet(doc, "组员自己领了个人任务 → 显示忙碌（可以和小组任务同时存在）")
    add_bullet(doc, "大组人多时，前端按每块最多 8 人拆成续页卡（角标 1/N），方便轮播看完，人数口径不变")

    add_heading(doc, "7.3 单工人", 2)
    add_para(
        doc,
        "单工人列表和小组是两套。入了组的人，只要还有个人任务，仍可能出现在单工人里。"
        "单工人闲忙只看个人任务，不看小组任务。",
    )

    add_heading(doc, "7.4 顶部忙 / 待命 / 闲", 2)
    add_bullet(doc, "忙：小组里标忙的人 + 单工人里标忙的人（按人名去重）")
    add_bullet(doc, "待命：组内待命人数")
    add_bullet(doc, "闲：总人数减去忙、再减去待命（总人数是两边出现过的人名去重）")

    # 8 大屏怎么排工单
    add_heading(doc, "八、大屏工单怎么一页排下来？", 1)
    add_para(
        doc,
        "工单列表按「高度格子」装箱，一页固定 6 格，卡片会按格子把高度均分，尽量铺满、少留白。",
    )
    add_table(
        doc,
        ["情况", "占几格", "大概一页几单"],
        [
            ["全是串行工单", "每单 1 格", "6 单"],
            ["有并行工单（默认只看前 4 道）", "并行那张占 2 格", "常见 5 单（1 并行 + 4 串行）"],
            [
                "并行点了「展示全部」",
                "先占 2 格，每多露出一道并行工序再 +1 格",
                "同页其它工单会变少",
            ],
        ],
    )
    add_bullet(doc, "并行区有「展示全部 / 收起」按钮")
    add_bullet(doc, "同一条路线里，下料、合套这类串行小盒子按正常高度，不会被并行虚线框拉高")
    add_bullet(doc, "工人区大组仍按每块最多 8 人拆续页，和工单区轮播秒数共用")

    # 9
    add_heading(doc, "九、联调最容易搞错的几件事", 1)
    add_table(
        doc,
        ["别踩坑", "正确理解"],
        [
            ["进行中只含「执行中」", "已下发、暂停也算进行中"],
            ["切周期会改进行中数字", "不会。周期只动「已完成」和「报工」"],
            ["报工数量必须凑满排产才算完成", "一次报工就算工序完成"],
            ["下发后直接执行中，算「待领取」", "算「在制」"],
            ["筛了车间，无工单号的散报工还算", "类型/中心有筛选时，散报工不算"],
            ["组内待命 = 空闲", "待命 ≠ 空闲；人还挂在组长接的小组活上"],
            ["人在小组里就不会出现在单工人", "有个人任务时，两边都可能出现"],
            ["一页条数随屏幕随便变", "工单区固定 6 格装箱；展开并行会占更多格"],
        ],
    )

    # 10
    add_heading(doc, "十、和开发对接时怎么说", 1)
    add_para(
        doc,
        "技术细节、字段名、建议接口结构，仍以仓库里的"
        "《工单监管看板-取数与交互说明.md》为准。"
        "原则一句话：以后换真实接口可以，业务口径不要改。",
    )
    add_para(doc, "日常使用相关：", bold=True)
    add_bullet(doc, "全屏入口路径：/board/work-order-monitor/screen")
    add_bullet(doc, "轮播开关和秒数会记在本机浏览器里，换电脑不会跟着走")
    add_bullet(doc, "大组 8 人分块、工单 6 格装箱与并行展开，都是前端展示，后端按整组/整单返回即可")

    add_heading(doc, "十一、一句话总结", 1)
    add_para(
        doc,
        "看板就是把「此刻有哪些单在跑、工序走到哪、谁在干活」摊开在大屏上；"
        "周期只管已经办完的单和报工成绩单；进行中永远看当下；"
        "工序报一次就算完；组内待命不是闲着；"
        "一页按 6 格排，并行默认占 2 格，展开全部会多占格、同页少看几单。",
    )

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    build()
