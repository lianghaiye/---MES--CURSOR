#!/usr/bin/env python3
"""生成《双物料单位-单物料双计量-vs-SKU拆解-决策说明.docx》"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "双物料单位-单物料双计量-vs-SKU拆解-决策说明.docx"


def font(run, size=11, bold=False, color=None, name="微软雅黑"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(
            r,
            size={1: 18, 2: 14, 3: 12}.get(level, 12),
            bold=True,
            color=(0x1F, 0x4E, 0x79) if level == 1 else (0x2E, 0x75, 0xB6),
        )
    return p


def P(doc, text, *, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def B(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    font(r)
    p.paragraph_format.space_after = Pt(3)
    return p


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    font(r1, size=10, bold=True, color=(0x1F, 0x4E, 0x79))
    for line in body.split("\n"):
        p = cell.add_paragraph()
        r = p.add_run(line)
        font(r, size=10)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        font(r, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            font(r, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("双物料单位建模决策说明")
    font(r, size=20, bold=True, color=(0x1F, 0x4E, 0x79))
    title.paragraph_format.space_after = Pt(12)

    callout(
        doc,
        "结论（采用）",
        "型材/板材类双单位（采购单位 ≠ 库存单位）采用「单物料 + 双计量」；\n"
        "不以拆 SKU 承载连续长度/面积。\n"
        "库存、BOM、下料只认一个库存料号；采购单位做换算与录入，不平行建存货 SKU。",
    )

    add_table(
        doc,
        ["项", "内容"],
        [
            ["适用范围", "下料、领料/发料出库、余料回库、库存台账、计划齐套占用"],
            ["适用物料", "钢管、棒料、板材等：采购按根/件/盒，库存按米/㎡（或重量）"],
            ["不适用", "规格/材质等可枚举差异（那是 SPU→SKU 变体，与双单位正交）"],
        ],
    )

    H(doc, "1. 两种方案定义", 2)
    add_table(
        doc,
        ["", "方案 A：单物料双计量（采用）", "方案 B：SKU 拆解（不采用）"],
        [
            [
                "主数据",
                "一个料号；采购单位 + 库存单位；批次带单件长度/面积",
                "拆成多个料号，如「按根 SKU」与「按米 SKU」，或「6米管 / 5米管 / 余料」各一码",
            ],
            [
                "库存账",
                "只认库存单位（米/㎡）；「根数」为辅助属性",
                "每个 SKU 各自余额；靠转换/虚拟出入库对齐",
            ],
            [
                "下料",
                "同料号：整出 → 实耗 → 余料回库",
                "跨 SKU：A 出 → B/C 入，或按长度换码",
            ],
        ],
    )

    H(doc, "2. 决策要点（为何不用 SKU 扛双单位）", 2)
    B(doc, "连续量 ≠ 可枚举规格：米/㎡/任意余料长度不能靠主数据枚举；强行拆 SKU → 编码爆炸或精度丢失。")
    B(doc, "实物同一性：一根管锯短后仍是同一物料；拆 SKU 把「数量变化」误做成「物料变化」。")
    B(doc, "账链路：跨 SKU 转换漏一笔即账实不符；同料号数量增减可闭环追溯。")
    B(doc, "计划/采购/库存三口径：一料号易统一；双 SKU 易双算或漏算缺口。")
    B(doc, "与变体正交：叶轮「规格×材质」用 SPU/SKU；钢管「根↔米」用 UOM，二者不要混用。")

    H(doc, "3. 反例：6 米整根领出，锯成耗用 5 米 + 余料 1 米", 2)
    P(
        doc,
        "前提：物料 WL-PIPE-Q235-50，采购单位=根，库存单位=米；原料仓有一批次，单件 6 米 × 1 根。",
    )

    H(doc, "3.1 方案 A（采用）— 记账", 3)
    add_table(
        doc,
        ["步骤", "动作", "库存变化（同一料号）"],
        [
            ["1", "领料出库（整出）原料仓 → 线边仓", "原料仓 −6 米；线边仓 +6 米（仍 1 根属性）"],
            ["2", "下料结算确认：实耗 5 米", "线边仓 −5 米（耗用）"],
            ["3", "余料 1 米回原料仓（或指定回仓）", "线边仓 −1 米；原料仓 +1 米（余料批次，长度=1）"],
        ],
    )
    P(doc, "结果：全程一个料号；台账净效果=原料仓少 5 米（投入生产）；批次可追溯出库单→结算单。")

    H(doc, "3.2 方案 B（不采用）— 典型记账", 3)
    add_table(
        doc,
        ["步骤", "动作", "库存变化（多料号）"],
        [
            ["1", "出「6米管」SKU，或「按根」转「按米」", "依赖转换单；漏转则两边账不平"],
            ["2", "耗用 5 米", "若只有「米」SKU，还要区分是否另建「下料耗用」逻辑"],
            [
                "3",
                "余料 1 米",
                "要么新建「1米管」SKU（长度一变一码），要么进笼统「余料」SKU（丢失长度，后续无法再拣选准确长度）",
            ],
        ],
    )
    P(doc, "结果：主数据膨胀或库存失真；追溯需跨多个 SKU 拼链路；结算难以标准化。")

    P(doc, "对照示意：", bold=True)
    P(doc, "方案 A：[同一料号]  6米 ──锯──► 耗5米 + 余1米（数量/批次属性变）")
    P(doc, "方案 B：[SKU-6米] ──────► [SKU-耗用?] + [SKU-1米 或 余料笼统码]（编码变）")

    H(doc, "4. 采用方案的边界约定（避免误解）", 2)
    add_table(
        doc,
        ["约定", "说明"],
        [
            ["库存账唯一口径", "账面数量 = 库存单位（米/㎡等）"],
            ["采购单位角色", "仅用于采购/入库录入与换算，不另建存货 SKU"],
            ["批次属性", "单件长度/面积、条码、货位；支持整出与余料回"],
            ["下料结算", "挂出库单；确认实耗与余料回库，不改料号"],
            ["SKU/变体", "仅用于规格、材质等离散维度，不用于长度分段"],
        ],
    )

    H(doc, "5. 一句话答复研发", 2)
    callout(
        doc,
        "对外口径",
        "双单位解决的是「同一存货的两种计量」，不是「两种存货」。\n"
        "用 SKU 拆根/米或按长度建码，短期单位看起来简单，长期主数据与对账成本不可控，"
        "且无法干净表达整出+余料回。\n"
        "库存、BOM、下料只认一个库存料号；采购单位做换算与录入，不平行建存货 SKU。",
    )

    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()
